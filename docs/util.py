from docx import Document as DocxDocument
import os
import time
from django.core.files.storage import default_storage


class Docx_Reader:
    def __init__(self, file_path, temp):
        self.docx_file = file_path
        self.temp = temp
        self.doc = DocxDocument(self.docx_file)

        self.tables = []

        self.variables = []

    def headers_text(self):
        paragraphs = []
        headings = []
        for para in self.doc.paragraphs:
            # Извлекаем заголовки
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split()[-1])  # Уровень заголовка (1, 2, 3)
                headings.append({
                    'text': para.text,
                    'level': level,
                    'id': f"heading-{len(headings)}"  # Уникальный ID для якоря
                })

            # Извлекаем текст
            paragraphs.append({
                'text': para.text,
                'style': para.style.name,
                'runs': [{'text': run.text, 'bold': run.bold, 'italic': run.italic, 'underline': run.underline} for run in
                         para.runs]
            })

        return headings, paragraphs
            # if self.temp:
            #     var = extract_variables(para.text)
            #     variables.append(var)


    def get_tables(self):
        # Извлекаем таблицы
        for table in self.doc.tables:
            rows = []
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = ''
                    for para in cell.paragraphs:
                        cell_text += para.text + '<br>'
                    cells.append(cell_text)
                rows.append(cells)
            self.tables.append(rows)
        return self.tables

    # clean_variables = clean_var(variables)

    def itog(self):
        heads, text = self.headers_text()
        context = {
            'document': self.doc,
            'paragraphs': text,
            'tables': self.get_tables(),
            'headings': heads,
            # 'page_obj': page_obj,
            # 'variables': clean_variables,
            # 'is_template': is_template,
        }
        return context


def make_file_unique(file_field, document):
    """
    Делает файл уникальным, если он совпадает с другими файлами документа.
    Переименовывает файл с timestamp суффиксом для избежания конфликтов.
    """
    if not file_field:
        return
    
    # Получаем относительный путь файла (без MEDIA_ROOT)
    file_path = file_field.name if hasattr(file_field, 'name') else str(file_field)
    
    conflicting_paths = set()
    
    # Собираем пути всех существующих файлов основного документа
    if document.file:
        conflicting_paths.add(document.file.name)
    if document.word_file:
        conflicting_paths.add(document.word_file.name)
    
    # Собираем пути всех файлов версий
    for version in document.versions.all():
        if version.file:
            conflicting_paths.add(version.file.name)
        if version.word_file:
            conflicting_paths.add(version.word_file.name)
    
    # Если есть конфликт — переименовываем
    if file_path in conflicting_paths:
        name, ext = os.path.splitext(file_path)
        timestamp = int(time.time() * 1000)  # миллисекунды для большей уникальности
        
        # Создаем новое имя файла
        new_name = f"{name}_{timestamp}{ext}"
        
        # Обновляем имя файла
        file_field.name = new_name
        
        print(f"DEBUG: File renamed from {file_path} to {new_name} to avoid conflict")